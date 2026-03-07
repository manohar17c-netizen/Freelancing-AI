from __future__ import annotations

import hashlib
import io
import re
import uuid
from typing import Dict, List, Tuple

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from pydantic import BaseModel, Field

from auth import auth_router
from database import (
    init_db,
    list_freelancer_profiles,
    load_freelancer_records,
    upsert_freelancer_record,
)
from ui_routes import ui_router

try:
    from sentence_transformers import SentenceTransformer
except ImportError:  # pragma: no cover
    SentenceTransformer = None

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover
    PdfReader = None

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None


app = FastAPI(title="Freelancing AI Matching API")
app.include_router(auth_router)
app.include_router(ui_router)


class InMemoryCollection:
    """A tiny in-memory substitute for vector DB behaviour."""

    def __init__(self) -> None:
        self._rows: List[Dict] = []

    def add(self, documents: List[str], embeddings: List[List[float]], ids: List[str]) -> None:
        for doc, emb, item_id in zip(documents, embeddings, ids):
            self._rows.append({"id": item_id, "document": doc, "embedding": emb})

    def query(self, query_embedding: List[float], n_results: int = 5) -> List[Dict]:
        scored = []
        for row in self._rows:
            score = cosine_similarity(query_embedding, row["embedding"])
            scored.append({"id": row["id"], "document": row["document"], "score": round(score, 4)})
        scored.sort(key=lambda item: item["score"], reverse=True)
        return scored[:n_results]


class JobPost(BaseModel):
    description: str = Field(..., min_length=10)
    required_skills: List[str] = Field(default_factory=list)
    min_experience_years: int = 0


def cosine_similarity(a: List[float], b: List[float]) -> float:
    dot = sum(x * y for x, y in zip(a, b))
    mag_a = sum(x * x for x in a) ** 0.5
    mag_b = sum(y * y for y in b) ** 0.5
    if mag_a == 0 or mag_b == 0:
        return 0.0
    return dot / (mag_a * mag_b)


def _fallback_embedding(text: str, dim: int = 32) -> List[float]:
    digest = hashlib.sha256(text.encode("utf-8")).digest()
    return [((digest[i % len(digest)] / 255.0) * 2) - 1 for i in range(dim)]


class EmbeddingService:
    def __init__(self) -> None:
        self.model = SentenceTransformer("all-MiniLM-L6-v2") if SentenceTransformer else None

    def encode(self, texts: List[str]) -> List[List[float]]:
        if self.model:
            vectors = self.model.encode(texts)
            return [vector.tolist() for vector in vectors]
        return [_fallback_embedding(text) for text in texts]


embedder = EmbeddingService()
resume_collection = InMemoryCollection()
resume_metadata: Dict[str, Dict] = {}


@app.on_event("startup")
def initialize_persistence() -> None:
    init_db()
    for record in load_freelancer_records():
        embedding = record.get("embedding") or embedder.encode([record["resume_text"]])[0]
        resume_collection.add(
            documents=[record["resume_text"]],
            embeddings=[embedding],
            ids=[record["resume_id"]],
        )
        resume_metadata[record["resume_id"]] = {
            "name": record["name"],
            "email": record["email"],
            "headline": record["headline"],
            "bio": record["bio"],
            "experience_years": max(int(record["experience_years"]), 0),
            "experience_months": max(min(int(record.get("experience_months", 0) or 0), 11), 0),
            "skills": sorted({skill.strip().lower() for skill in record["skills"] if isinstance(skill, str)}),
        }


@app.get("/")
def healthcheck() -> Dict[str, str]:
    return {"status": "ok", "service": "Freelancing AI Matching API", "ui": "/ui"}


@app.post("/freelancers/register")
async def register_freelancer(
    name: str = Form(...),
    email: str = Form(...),
    headline: str = Form(...),
    experience_years: int = Form(0),
    experience_months: int = Form(0),
    allow_empty_resume: bool = Form(False),
    skills: str = Form(""),
    bio: str = Form(""),
    file: UploadFile = File(...),
) -> Dict:
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded resume is empty.")

    text, extraction_warning = _extract_resume_text(file, content, allow_empty_resume=allow_empty_resume)

    resume_id = f"resume-{uuid.uuid4().hex[:8]}"
    embedding = embedder.encode([text])[0]
    resume_collection.add(documents=[text], embeddings=[embedding], ids=[resume_id])

    typed_skills = [skill.strip().lower() for skill in skills.split(",") if skill.strip()]
    inferred_skills = _extract_skills(text.lower())
    merged_skills = sorted(set(typed_skills) | set(inferred_skills))
    extracted_years, extracted_months = _extract_experience_from_text(text.lower()) if text.strip() else (0, 0)

    normalized_months = max(min(experience_months, 11), 0)
    total_input_months = max(experience_years, 0) * 12 + normalized_months
    total_extracted_months = max(extracted_years, 0) * 12 + max(min(extracted_months, 11), 0)
    final_total_months = max(total_input_months, total_extracted_months)

    resume_metadata[resume_id] = {
        "name": name,
        "email": email,
        "headline": headline,
        "bio": bio,
        "experience_years": final_total_months // 12,
        "experience_months": final_total_months % 12,
        "skills": merged_skills,
    }

    upsert_freelancer_record(
        resume_id=resume_id,
        name=name,
        email=email,
        headline=headline,
        bio=bio,
        experience_years=resume_metadata[resume_id]["experience_years"],
        experience_months=resume_metadata[resume_id]["experience_months"],
        skills=merged_skills,
        resume_text=text,
        embedding=embedding,
    )

    response = {
        "message": "Freelancer registered successfully",
        "resume_id": resume_id,
        "profile": {"name": name, "email": email, "headline": headline, "skills": merged_skills},
    }
    if extraction_warning:
        response["warning"] = extraction_warning
    return response


@app.post("/upload-resume")
async def upload_resume(file: UploadFile = File(...), allow_empty_resume: bool = Form(False)) -> Dict:
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded resume is empty.")

    text, extraction_warning = _extract_resume_text(file, content, allow_empty_resume=allow_empty_resume)

    resume_id = f"resume-{uuid.uuid4().hex[:8]}"
    embedding = embedder.encode([text])[0]
    resume_collection.add(documents=[text], embeddings=[embedding], ids=[resume_id])

    lowered = text.lower() if text.strip() else ""
    extracted_years, extracted_months = _extract_experience_from_text(lowered)
    resume_metadata[resume_id] = {
        "name": file.filename,
        "email": "",
        "headline": "Freelancer",
        "bio": "",
        "experience_years": extracted_years,
        "experience_months": extracted_months,
        "skills": _extract_skills(lowered),
    }

    upsert_freelancer_record(
        resume_id=resume_id,
        name=file.filename,
        email="",
        headline="Freelancer",
        bio="",
        experience_years=resume_metadata[resume_id]["experience_years"],
        experience_months=resume_metadata[resume_id]["experience_months"],
        skills=resume_metadata[resume_id]["skills"],
        resume_text=text,
        embedding=embedding,
    )

    response = {"message": "Resume stored successfully", "resume_id": resume_id}
    if extraction_warning:
        response["warning"] = extraction_warning
    return response


@app.post("/post-job")
def post_job(job: JobPost) -> Dict:
    if not resume_metadata:
        raise HTTPException(status_code=404, detail="No resumes found. Upload resumes first.")

    job_embedding = embedder.encode([job.description])[0]
    semantic_matches = resume_collection.query(query_embedding=job_embedding, n_results=10)

    ranked = []
    for result in semantic_matches:
        candidate = resume_metadata.get(result["id"], {"skills": [], "experience_years": 0, "experience_months": 0})
        candidate_total_months = max(candidate["experience_years"], 0) * 12 + max(candidate.get("experience_months", 0), 0)
        required_total_months = max(job.min_experience_years, 0) * 12
        experience_score = min(candidate_total_months / max(required_total_months, 1), 1.0)

        required = {skill.lower() for skill in job.required_skills}
        available = set(candidate["skills"])
        skill_score = len(required & available) / max(len(required), 1)

        final_score = round((0.5 * skill_score) + (0.3 * experience_score) + (0.2 * result["score"]), 4)
        ranked.append(
            {
                "resume_id": result["id"],
                "name": candidate.get("name", result["id"]),
                "headline": candidate.get("headline", "Freelancer"),
                "semantic_score": result["score"],
                "skill_score": round(skill_score, 4),
                "experience_score": round(experience_score, 4),
                "experience_years": candidate.get("experience_years", 0),
                "experience_months": candidate.get("experience_months", 0),
                "final_score": final_score,
            }
        )

    ranked.sort(key=lambda item: item["final_score"], reverse=True)
    return {"top_matches": ranked[:5]}


@app.get("/freelancers")
def get_freelancers(limit: int = 100) -> Dict[str, List[Dict]]:
    profiles = list_freelancer_profiles(limit)
    return {"count": len(profiles), "freelancers": profiles}


def _extract_resume_text(file: UploadFile, content: bytes, allow_empty_resume: bool = False) -> Tuple[str, str]:
    filename = (file.filename or "").lower()
    extension = "." + filename.rsplit(".", 1)[-1] if "." in filename else ""

    def _empty_text_fallback(detail: str) -> Tuple[str, str]:
        if allow_empty_resume:
            fallback_text = f"[UNPARSEABLE_RESUME:{extension or 'unknown'}] {detail}"
            return fallback_text, detail
        raise HTTPException(status_code=400, detail=detail)

    if filename.endswith(".txt"):
        text = content.decode("utf-8", errors="ignore").strip()
        if text:
            return text, ""
        return _empty_text_fallback("TXT has no readable text content.")

    if filename.endswith(".pdf"):
        if PdfReader is None:
            raise HTTPException(status_code=500, detail="PDF support requires pypdf. Install dependency and retry.")
        reader = PdfReader(io.BytesIO(content))
        text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
        if text:
            return text, ""
        return _empty_text_fallback(
            "PDF text extraction returned empty content. The PDF is likely scanned/image-only or protected."
        )

    if filename.endswith(".docx"):
        if Document is None:
            raise HTTPException(status_code=500, detail="DOCX support requires python-docx. Install dependency and retry.")
        document = Document(io.BytesIO(content))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text and paragraph.text.strip()]
        table_cells = []
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        table_cells.append(cell_text)
        text = "\n".join(paragraphs + table_cells).strip()
        if text:
            return text, ""
        return _empty_text_fallback("DOCX has no readable text in paragraphs or tables.")

    if filename.endswith(".doc"):
        # Legacy .doc is binary and not reliably parseable without external converters.
        decoded = content.decode("utf-8", errors="ignore").strip()
        if decoded:
            return decoded, "Legacy DOC parsed with fallback decoder. Convert to DOCX for better accuracy."
        return _empty_text_fallback(
            "Legacy DOC could not be parsed reliably. Convert the file to DOCX and upload again."
        )

    raise HTTPException(status_code=400, detail="Unsupported resume format. Upload TXT, PDF, DOCX, or DOC.")


def _extract_experience_from_text(resume_text_lower: str) -> tuple[int, int]:
    years = 0
    months = 0

    year_match = re.search(r"(\d{1,2})\s*\+?\s*(?:years?|yrs?)", resume_text_lower)
    month_match = re.search(r"(\d{1,2})\s*\+?\s*(?:months?|mos?)", resume_text_lower)

    if year_match:
        years = max(0, min(int(year_match.group(1)), 50))
    if month_match:
        months = max(0, min(int(month_match.group(1)), 11))

    if years == 0 and months == 0:
        for token in resume_text_lower.split():
            if token.isdigit():
                value = int(token)
                if 0 < value < 51:
                    years = value
                    break

    return years, months


def _extract_skills(resume_text_lower: str) -> List[str]:
    common_skills = {
        "python",
        "fastapi",
        "django",
        "flask",
        "react",
        "next.js",
        "node",
        "aws",
        "postgresql",
        "docker",
        "kubernetes",
    }
    return sorted(skill for skill in common_skills if skill in resume_text_lower)
